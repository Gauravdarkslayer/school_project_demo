from django.shortcuts import render
from django.http import HttpResponse
from .models import organization_test_create
from django.contrib import messages
# Create your views here.

def show_doc_upload_page(request):
    return render(request,'doc_upload_1.html')


def doc_processing(request):
    file = request.FILES["mydocument"]
    # print(request.FILES)
    print("this is file object",file)
    # current_user_email = request.session["email"]
    import docx
    try:
        doc = docx.Document('file.docx')
        print("success")
    except Exception as e:
        print(e)
    
    doc = docx.Document(file)
    print("this works")
    questionnaire={}
    explanation={}
    diagram={}
    fill_blanks = {}
    i=1
    number_of_iterations = 0
    for content in doc.paragraphs:
        if ("Q"+str(i) or "QUESTION"+str(i)) in content.text.upper():
            print(content.text)
            # questionnaire["question"+str(i)] = content.text
            questionnaire[content.text] =[]
            i=i+1
            j=1
            for option in doc.paragraphs:
                # print("this is option",option.text.upper())
                if "OPTION"+" "+str(j) in option.text.upper():
                    print("check")
                    questionnaire[content.text].append(option.text)
                    j= j+1
                if "ANSWER" in option.text.upper():
                    print("get answer")
                    questionnaire[content.text].append(option.text)
                    break 
        elif ('fill in the blanks') in content.text.lower():
            print("I am inside Fill in the Blanks")
            fill_blanks[content.text] = []
            next_index = number_of_iterations+1
            
            for i in range(next_index, len(doc.paragraphs)):
                if "_" in doc.paragraphs[i].text:
                    fill_blanks[content.text].append(doc.paragraphs[i].text)
                elif " " == doc.paragraphs[i].text:
                    continue
                else:
                    break
        elif("explain") in content.text.lower():
            explanation[content.text]=[]
            next_index = number_of_iterations+1
            for i in range(next_index, len(doc.paragraphs)):
                if "Write" in doc.paragraphs[i].text:
                    explanation[content.text].append(doc.paragraphs[i].text)
                elif " " == doc.paragraphs[i].text:
                    continue
                else:
                    break
        elif("draw") in content.text.lower():
            diagram[content.text]=[]
            next_index = number_of_iterations+1
            for i in range(next_index, len(doc.paragraphs)):
                if "diagram" in doc.paragraphs[i].text:
                    diagram[content.text].append(doc.paragraphs[i].text)
                elif " " == doc.paragraphs[i].text:
                    continue
                else:
                    break

            print("Explanation",explanation)
            print("Diagram",diagram)
            print("Data of Fill in the blanks is ", fill_blanks)
        number_of_iterations += 1
    print(questionnaire)

    return render(request,"mcq.html",{'context':questionnaire, 'fill_blanks':fill_blanks,"explain":explanation,"diagram":diagram})


def preview(request):
    if request.method == "POST":
        try:
            name = request.POST["organization_name"]
            description = request.POST["description"]
            logo = request.FILES["logo"]
            test_name = request.POST["test_name"]
            college_name = request.POST["college_name"]
            cam_micro = request.POST["cam_micro"]
            test_duration = request.POST["test_duration"]

            organization_test_create.objects.create(
                name = name,
                description=description,
                logo = logo,
                test_name = test_name,
                college_name = college_name,
                cam_micro = cam_micro,
                test_duration = test_duration
            )
            context = {
                "name":name,
                "description":description,
                "logo":logo,
                "test_name":test_name,
                "college_name":college_name,
                "cam_micro":cam_micro,
                "test_duration":test_duration
            }
            messages.success(request,"Test created successfully")
            return render(request,"form.html",context=context)
        except Exception as e:
            messages.error(request,e)
            return render(request,"form.html")
    else:
        return render(request,"form.html")
