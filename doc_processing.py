import docx
def getText(doc):
    # doc = docx.Document(filename)
    fullText = []
    for para in doc.paragraphs:
        fullText.append(para.text)
    return '\n'.join(fullText)

doc = docx.Document('test_doc.docx')
# print(len(doc.paragraphs))
# print(doc.paragraphs[0].text)
# print(getText(doc))
# actual_content = getText(doc)
questionnaire={}
i=1
# print(getText(doc))

# working code
for content in doc.paragraphs:
    if "Q"+str(i) in content.text.upper():
        
        # questionnaire["question"+str(i)] = content.text
        questionnaire[content.text] =[]
        i=i+1
        j=1
        for option in doc.paragraphs:
            # print("this is option",option.text.upper())
            if "OPTION"+" "+str(j) in option.text.upper():
                print("check")
                questionnaire[content.text].append(option.text)
                j= j+1
print(questionnaire) 

# for content in doc.paragraphs:
#     if "Q"+str(i) in content.text.upper():
        
#         # questionnaire["question"+str(i)] = content.text
#         questionnaire[content.text] =[]
#         i=i+1
#         j=1
#         question
#     elif "OPTION"+" "+str(j) in content.text.upper():
#         questionnaire[content.text].append(option.text)
#         j= j+1

        # for option in doc.paragraphs:
        #     # print("this is option",option.text.upper())
        #     if "OPTION"+" "+str(j) in option.text.upper():
        #         print("check")
        #         questionnaire[content.text].append(option.text)
        #         j= j+1
"""
{
    "question1":[option1,option2,option3,option4,answer],
    "question2":[option1,option2,option3,option4,answer],
    
}

"""